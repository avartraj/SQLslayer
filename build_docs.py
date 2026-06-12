"""
build_docs.py — Convert the Markdown documentation to styled Word (.docx).

Both documents are rendered by the SAME function with the SAME theme, so they
share identical formatting (fonts, heading colours, tables, code blocks).

Usage:  python build_docs.py
Outputs: SQLSlayer_Documentation.docx, VulnShop_Documentation.docx
"""
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── theme ────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x4E, 0x79)
BLUE   = RGBColor(0x2E, 0x74, 0xB5)
GREY   = RGBColor(0x59, 0x59, 0x59)
CODECL = RGBColor(0xA3, 0x14, 0x45)   # inline-code colour
BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
CODE_SHADE = "F2F2F2"
TABLE_STYLE = "Light Grid Accent 1"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _shade(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _add_inline(paragraph, text):
    """Add a run sequence honouring **bold** and `code` markup."""
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = paragraph.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = paragraph.add_run(piece[1:-1])
            r.font.name = CODE_FONT; r.font.size = Pt(9.5); r.font.color.rgb = CODECL
        else:
            paragraph.add_run(piece)


def _strip_inline(text):
    return text.replace("**", "").replace("`", "")


def _set_base_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)


def _heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    if level == 1:
        r = p.add_run(_strip_inline(text)); r.bold = True
        r.font.size = Pt(22); r.font.color.rgb = NAVY
    elif level == 2:
        r = p.add_run(_strip_inline(text)); r.bold = True
        r.font.size = Pt(15); r.font.color.rgb = NAVY
        _bottom_border(p)
    else:
        r = p.add_run(_strip_inline(text)); r.bold = True
        r.font.size = Pt(12.5); r.font.color.rgb = BLUE
    return p


def _code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.15)
        _shade(p, CODE_SHADE)
        r = p.add_run(ln if ln else " ")
        r.font.name = CODE_FONT; r.font.size = Pt(9)


def _table(doc, rows):
    # drop separator rows (---), detect empty header
    rows = [r for r in rows if not all(re.fullmatch(r"\s*:?-{2,}:?\s*", c or "") for c in r)]
    if not rows:
        return
    header_empty = all(not c.strip() for c in rows[0])
    body = rows[1:] if header_empty else rows
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncols)
    try:
        t.style = TABLE_STYLE
    except KeyError:
        t.style = "Table Grid"
    has_header = not header_empty
    for i, row in enumerate(body if header_empty else rows):
        cells = t.add_row().cells
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            para = cells[j].paragraphs[0]
            _add_inline(para, val.strip())
            if (has_header and i == 0) or (header_empty and j == 0):
                for run in para.runs:
                    run.bold = True


def render(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    _set_base_style(doc)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        # code fence
        if line.lstrip().startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].lstrip().startswith("```"):
                block.append(lines[i]); i += 1
            i += 1
            _code_block(doc, block)
            continue

        # table (consecutive | ... | lines)
        if line.strip().startswith("|") and line.strip().endswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c for c in lines[i].strip().strip("|").split("|")]
                tbl.append(cells); i += 1
            _table(doc, tbl)
            doc.add_paragraph()
            continue

        stripped = line.strip()
        if stripped.startswith("# "):
            _heading(doc, stripped[2:], 1)
        elif stripped.startswith("### "):
            _heading(doc, stripped[4:], 3)
        elif stripped.startswith("## "):
            _heading(doc, stripped[3:], 2)
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph(); _bottom_border(p)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            _add_inline(p, stripped[2:])
            for r in p.runs:
                r.italic = True; r.font.color.rgb = GREY
        elif re.match(r"^[-*] ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif not stripped:
            pass
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

        i += 1

    doc.save(docx_path)
    print(f"  wrote {docx_path}")


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    pairs = [
        ("SQLSlayer_Documentation.md", "SQLSlayer_Documentation.docx"),
        ("VulnShop_Documentation.md", "VulnShop_Documentation.docx"),
    ]
    print("Building Word documents (shared theme):")
    for md, dx in pairs:
        render(os.path.join(here, md), os.path.join(here, dx))
    print("Done.")
